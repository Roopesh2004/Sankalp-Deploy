#!/usr/bin/env python3
"""
Test all PDF conversion methods for the certificate service
"""

import os
import sys
import subprocess
from docx import Document

# Add backend directory to path
sys.path.insert(0, 'backend')

def test_docx2pdf():
    """Test docx2pdf conversion"""
    print("Testing docx2pdf conversion...")
    
    try:
        from docx2pdf import convert
        print("✓ docx2pdf library available")
        
        # Create test document
        doc = Document()
        doc.add_heading('Test Certificate', 0)
        doc.add_paragraph('This is a test for docx2pdf conversion.')
        
        test_docx = "test_docx2pdf.docx"
        test_pdf = "test_docx2pdf.pdf"
        
        doc.save(test_docx)
        
        # Try conversion
        try:
            convert(test_docx, test_pdf)
            if os.path.exists(test_pdf):
                print("✓ docx2pdf conversion successful")
                os.remove(test_pdf)
                success = True
            else:
                print("✗ docx2pdf conversion failed - no PDF created")
                success = False
        except Exception as e:
            print(f"✗ docx2pdf conversion failed: {e}")
            success = False
        
        # Cleanup
        if os.path.exists(test_docx):
            os.remove(test_docx)
        
        return success
        
    except ImportError:
        print("✗ docx2pdf library not available")
        return False

def test_libreoffice():
    """Test LibreOffice conversion"""
    print("\nTesting LibreOffice conversion...")
    
    # Import the function from our service
    try:
        from certificate_service import convert_with_libreoffice
        
        # Create test document
        doc = Document()
        doc.add_heading('Test Certificate', 0)
        doc.add_paragraph('This is a test for LibreOffice conversion.')
        
        test_docx = "test_libreoffice.docx"
        test_pdf = "test_libreoffice.pdf"
        
        doc.save(test_docx)
        
        # Try conversion
        success = convert_with_libreoffice(test_docx, test_pdf)
        
        if success and os.path.exists(test_pdf):
            print("✓ LibreOffice conversion successful")
            os.remove(test_pdf)
        else:
            print("✗ LibreOffice conversion failed")
            success = False
        
        # Cleanup
        if os.path.exists(test_docx):
            os.remove(test_docx)
        
        return success
        
    except Exception as e:
        print(f"✗ LibreOffice test failed: {e}")
        return False

def test_reportlab():
    """Test reportlab fallback conversion"""
    print("\nTesting reportlab fallback conversion...")
    
    try:
        from certificate_service import convert_with_reportlab
        
        test_docx = "test_reportlab.docx"  # Not actually used by reportlab
        test_pdf = "test_reportlab.pdf"
        
        # Test data
        name = "John Doe"
        domain = "Web Development"
        start_date = "January 1, 2024"
        end_date = "March 31, 2024"
        gender = "male"
        
        # Try conversion
        success = convert_with_reportlab(test_docx, test_pdf, name, domain, start_date, end_date, gender)
        
        if success and os.path.exists(test_pdf):
            print("✓ reportlab conversion successful")
            file_size = os.path.getsize(test_pdf)
            print(f"  PDF size: {file_size} bytes")
            os.remove(test_pdf)
        else:
            print("✗ reportlab conversion failed")
            success = False
        
        return success
        
    except Exception as e:
        print(f"✗ reportlab test failed: {e}")
        return False

def test_full_service():
    """Test the full certificate generation service"""
    print("\nTesting full certificate generation service...")
    
    try:
        from certificate_service import generate_certificate
        
        # Change to backend directory
        original_cwd = os.getcwd()
        os.chdir('backend')
        
        try:
            # Test data
            name = "Jane Smith"
            domain = "Data Science"
            start_date = "February 1, 2024"
            end_date = "April 30, 2024"
            gender = "female"
            
            # Generate certificate
            pdf_path = generate_certificate(name, domain, start_date, end_date, gender)
            
            if pdf_path and os.path.exists(pdf_path):
                print("✓ Full service test successful")
                file_size = os.path.getsize(pdf_path)
                print(f"  Generated PDF: {pdf_path} ({file_size} bytes)")
                os.remove(pdf_path)
                success = True
            else:
                print("✗ Full service test failed - no PDF generated")
                success = False
                
        finally:
            os.chdir(original_cwd)
        
        return success
        
    except Exception as e:
        print(f"✗ Full service test failed: {e}")
        return False

def main():
    """Main test function"""
    print("Certificate Service PDF Conversion Test")
    print("=" * 50)
    
    results = {}
    
    # Test individual conversion methods
    results['docx2pdf'] = test_docx2pdf()
    results['libreoffice'] = test_libreoffice()
    results['reportlab'] = test_reportlab()
    results['full_service'] = test_full_service()
    
    # Summary
    print("\n" + "=" * 50)
    print("TEST RESULTS SUMMARY")
    print("=" * 50)
    
    for method, success in results.items():
        status = "✓ PASS" if success else "✗ FAIL"
        print(f"{method:15} : {status}")
    
    # Overall assessment
    working_methods = sum(results.values())
    total_methods = len(results)
    
    print(f"\nWorking methods: {working_methods}/{total_methods}")
    
    if results['full_service']:
        print("\n🎉 Certificate service is working!")
    elif working_methods > 0:
        print("\n⚠️  Certificate service has some working conversion methods")
    else:
        print("\n❌ Certificate service is not working - no conversion methods available")
    
    return results['full_service']

if __name__ == "__main__":
    success = main()
    sys.exit(0 if success else 1)
