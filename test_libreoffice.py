#!/usr/bin/env python3
"""
Test LibreOffice PDF conversion functionality
"""

import subprocess
import os
import sys
from docx import Document

def test_libreoffice_availability():
    """Test if LibreOffice is available and working"""
    print("Testing LibreOffice availability...")
    
    libreoffice_commands = [
        'libreoffice',
        'soffice',
        '/usr/bin/libreoffice',
        '/usr/bin/soffice',
        '/opt/libreoffice/program/soffice'
    ]
    
    for cmd in libreoffice_commands:
        try:
            print(f"Trying command: {cmd}")
            result = subprocess.run([cmd, '--version'], 
                                  capture_output=True, text=True, timeout=10)
            if result.returncode == 0:
                print(f"✓ LibreOffice found: {cmd}")
                print(f"  Version: {result.stdout.strip()}")
                return cmd
            else:
                print(f"  Command failed with return code: {result.returncode}")
                if result.stderr:
                    print(f"  Error: {result.stderr}")
        except subprocess.TimeoutExpired:
            print(f"  Command timed out: {cmd}")
        except FileNotFoundError:
            print(f"  Command not found: {cmd}")
    
    print("✗ No working LibreOffice installation found")
    return None

def create_test_docx():
    """Create a simple test DOCX file"""
    print("Creating test DOCX file...")
    
    doc = Document()
    doc.add_heading('Test Certificate', 0)
    doc.add_paragraph('This is a test certificate for {{Name}}.')
    doc.add_paragraph('Domain: {{Domain}}')
    doc.add_paragraph('Start Date: {{Start Date}}')
    doc.add_paragraph('End Date: {{End Date}}')
    
    test_docx = "test_conversion.docx"
    doc.save(test_docx)
    print(f"✓ Test DOCX created: {test_docx}")
    return test_docx

def test_conversion(libreoffice_cmd, input_docx):
    """Test the actual conversion"""
    print(f"Testing conversion with {libreoffice_cmd}...")
    
    output_pdf = "test_conversion.pdf"
    
    try:
        # Remove existing PDF if it exists
        if os.path.exists(output_pdf):
            os.remove(output_pdf)
        
        # Run conversion
        result = subprocess.run([
            libreoffice_cmd,
            '--headless',
            '--convert-to', 'pdf',
            '--outdir', '.',
            input_docx
        ], capture_output=True, text=True, timeout=30)
        
        print(f"Conversion result: return_code={result.returncode}")
        if result.stdout:
            print(f"stdout: {result.stdout}")
        if result.stderr:
            print(f"stderr: {result.stderr}")
        
        # Check if PDF was created
        if os.path.exists(output_pdf):
            file_size = os.path.getsize(output_pdf)
            print(f"✓ PDF created successfully: {output_pdf} ({file_size} bytes)")
            return True
        else:
            print(f"✗ PDF not created: {output_pdf}")
            print(f"Files in current directory: {os.listdir('.')}")
            return False
            
    except subprocess.TimeoutExpired:
        print("✗ Conversion timed out")
        return False
    except Exception as e:
        print(f"✗ Conversion failed with error: {e}")
        return False

def cleanup():
    """Clean up test files"""
    test_files = ["test_conversion.docx", "test_conversion.pdf"]
    for file in test_files:
        if os.path.exists(file):
            os.remove(file)
            print(f"Cleaned up: {file}")

def main():
    """Main test function"""
    print("LibreOffice PDF Conversion Test")
    print("=" * 40)
    
    try:
        # Test LibreOffice availability
        libreoffice_cmd = test_libreoffice_availability()
        if not libreoffice_cmd:
            print("\n❌ LibreOffice is not available. PDF conversion will not work.")
            return False
        
        print()
        
        # Create test DOCX
        test_docx = create_test_docx()
        
        print()
        
        # Test conversion
        success = test_conversion(libreoffice_cmd, test_docx)
        
        print("\n" + "=" * 40)
        if success:
            print("🎉 LibreOffice PDF conversion is working!")
        else:
            print("❌ LibreOffice PDF conversion failed")
        
        return success
        
    except KeyboardInterrupt:
        print("\nTest interrupted by user")
        return False
    finally:
        cleanup()

if __name__ == "__main__":
    success = main()
    sys.exit(0 if success else 1)
