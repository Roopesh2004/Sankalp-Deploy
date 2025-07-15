from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
from docx import Document
from datetime import datetime
import os
import uuid
import subprocess
import platform
import logging
import sys

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.StreamHandler(sys.stdout)
    ]
)
logger = logging.getLogger(__name__)

# Try to import docx2pdf, but handle the case where it might not work
try:
    from docx2pdf import convert
    DOCX2PDF_AVAILABLE = True
    logger.info("docx2pdf library loaded successfully")
except ImportError:
    DOCX2PDF_AVAILABLE = False
    logger.warning("docx2pdf not available. PDF conversion may not work.")

# Try to import reportlab for fallback PDF generation
try:
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    REPORTLAB_AVAILABLE = True
    logger.info("reportlab library loaded successfully")
except ImportError:
    REPORTLAB_AVAILABLE = False
    logger.warning("reportlab not available. Fallback PDF generation may not work.")

app = Flask(__name__)
CORS(app)  # Enable CORS for all routes

def convert_with_libreoffice(input_docx, output_pdf):
    """
    Alternative PDF conversion using LibreOffice headless mode
    Returns True if successful, False otherwise
    """
    try:
        logger.info(f"Attempting LibreOffice conversion: {input_docx} -> {output_pdf}")

        # Try different LibreOffice executable names
        libreoffice_commands = [
            'libreoffice',
            'soffice',
            '/usr/bin/libreoffice',
            '/usr/bin/soffice',
            '/opt/libreoffice/program/soffice'
        ]

        for cmd in libreoffice_commands:
            try:
                logger.info(f"Trying LibreOffice command: {cmd}")

                # Run LibreOffice in headless mode to convert DOCX to PDF
                result = subprocess.run([
                    cmd,
                    '--headless',
                    '--convert-to', 'pdf',
                    '--outdir', os.path.dirname(output_pdf) or '.',
                    input_docx
                ], capture_output=True, text=True, timeout=30)

                logger.info(f"LibreOffice command result: return_code={result.returncode}")
                if result.stdout:
                    logger.info(f"LibreOffice stdout: {result.stdout}")
                if result.stderr:
                    logger.warning(f"LibreOffice stderr: {result.stderr}")

                if result.returncode == 0:
                    # LibreOffice creates PDF with same name as input but .pdf extension
                    expected_pdf = os.path.splitext(input_docx)[0] + '.pdf'
                    logger.info(f"Looking for generated PDF: {expected_pdf}")

                    if os.path.exists(expected_pdf):
                        logger.info(f"PDF generated successfully: {expected_pdf}")
                        # Rename to desired output name if different
                        if expected_pdf != output_pdf:
                            logger.info(f"Renaming {expected_pdf} to {output_pdf}")
                            os.rename(expected_pdf, output_pdf)
                        return True
                    else:
                        logger.warning(f"Expected PDF not found: {expected_pdf}")
                        # List files in current directory for debugging
                        current_files = os.listdir('.')
                        logger.info(f"Files in current directory: {current_files}")

            except subprocess.TimeoutExpired:
                logger.warning(f"LibreOffice command timed out: {cmd}")
                continue
            except FileNotFoundError:
                logger.warning(f"LibreOffice command not found: {cmd}")
                continue

        logger.error("All LibreOffice commands failed")
        return False

    except Exception as e:
        logger.error(f"LibreOffice conversion error: {e}")
        return False

def convert_with_reportlab(input_docx, output_pdf, name, domain, start_date, end_date, gender):
    """
    Fallback PDF generation using reportlab
    Creates a simple certificate PDF when other methods fail
    """
    if not REPORTLAB_AVAILABLE:
        logger.error("reportlab not available for fallback conversion")
        return False

    try:
        logger.info(f"Creating fallback PDF with reportlab: {output_pdf}")

        # Create PDF document
        doc = SimpleDocTemplate(output_pdf, pagesize=A4)
        styles = getSampleStyleSheet()
        story = []

        # Title
        title_style = styles['Title']
        title = Paragraph("CERTIFICATE OF COMPLETION", title_style)
        story.append(title)
        story.append(Spacer(1, 20))

        # Content
        normal_style = styles['Normal']

        # Determine pronouns
        pronouns = {"male": ("he", "him"), "female": ("she", "her"), "other": ("they", "them")}
        he_she, him_her = pronouns.get(gender.lower(), ("they", "them"))

        # Certificate text
        content = f"""
        <para align="center">
        This is to certify that<br/><br/>
        <b>{name}</b><br/><br/>
        has successfully completed the training program in<br/><br/>
        <b>{domain}</b><br/><br/>
        from {start_date} to {end_date}.<br/><br/>
        {he_she.capitalize()} has demonstrated proficiency in the subject matter
        and is hereby awarded this certificate.<br/><br/>
        Issued on: {datetime.today().strftime('%B %d, %Y')}
        </para>
        """

        para = Paragraph(content, normal_style)
        story.append(para)

        # Build PDF
        doc.build(story)

        logger.info(f"Fallback PDF created successfully: {output_pdf}")
        return True

    except Exception as e:
        logger.error(f"reportlab conversion error: {e}")
        return False

def generate_certificate(name, domain, start_date, end_date, gender):
    print("Name: ",name)
    """
    Generate a certificate with the provided details
    Returns the path to the generated PDF file
    """
    try:
        # === Pronoun mapping ===
        pronouns = {"male": ("he", "him"), "female": ("she", "her"), "other": ("they", "them")}
        he_she, him_her = pronouns.get(gender.lower(), ("they", "them"))

        issued_date = datetime.today().strftime('%B %d, %Y')

        # === Load and fill Word document ===
        template_path = "SpectoV_Cert.docx"
        if not os.path.exists(template_path):
            raise FileNotFoundError(f"Certificate template not found: {template_path}")
            
        doc = Document(template_path)

        # Replace placeholders in all paragraphs
        for para in doc.paragraphs:
            if "{{Domain}}" in para.text:
                para.text = para.text.replace("{{Domain}}", domain)
            if "{{Start Date}}" in para.text:
                para.text = para.text.replace("{{Start Date}}", start_date)
            if "{{End Date}}" in para.text:
                para.text = para.text.replace("{{End Date}}", end_date)
            if "{{he/she/they}}" in para.text:
                para.text = para.text.replace("{{he/she/they}}", he_she)
            if "{{him/her/them}}" in para.text:
                para.text = para.text.replace("{{him/her/them}}", him_her)
            if "{{Name}}" in para.text:
                para.text = para.text.replace("{{Name}}", name)
            if "ISSUED DATE :" in para.text:
                para.text = para.text.replace("ISSUED DATE :", f"ISSUED DATE : {issued_date}")

        # Also check tables for placeholders
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if "{{Domain}}" in para.text:
                            para.text = para.text.replace("{{Domain}}", domain)
                        if "{{Start Date}}" in para.text:
                            para.text = para.text.replace("{{Start Date}}", start_date)
                        if "{{End Date}}" in para.text:
                            para.text = para.text.replace("{{End Date}}", end_date)
                        if "{{he/she/they}}" in para.text:
                            para.text = para.text.replace("{{he/she/they}}", he_she)
                        if "{{him/her/them}}" in para.text:
                            para.text = para.text.replace("{{him/her/them}}", him_her)
                        if "{{Name}}" in para.text:
                            para.text = para.text.replace("{{Name}}", name)
                        if "ISSUED DATE :" in para.text:
                            para.text = para.text.replace("ISSUED DATE :", f"ISSUED DATE : {issued_date}")

        # === Generate unique filenames ===
        unique_id = str(uuid.uuid4())
        output_docx = f"temp_certificate_{unique_id}.docx"
        output_pdf = f"certificate_{unique_id}.pdf"
        
        # === Save DOCX ===
        doc.save(output_docx)

        # === Convert to PDF ===
        logger.info(f"Starting PDF conversion. docx2pdf available: {DOCX2PDF_AVAILABLE}")

        conversion_successful = False

        if DOCX2PDF_AVAILABLE:
            try:
                logger.info("Attempting conversion with docx2pdf")
                convert(output_docx, output_pdf)
                logger.info("docx2pdf conversion successful")
                conversion_successful = True
            except Exception as e:
                logger.warning(f"docx2pdf conversion failed: {e}")

        # Try LibreOffice if docx2pdf failed or is not available
        if not conversion_successful:
            logger.info("Falling back to LibreOffice conversion")
            if convert_with_libreoffice(output_docx, output_pdf):
                logger.info("Successfully converted using LibreOffice")
                conversion_successful = True
            else:
                logger.warning("LibreOffice conversion failed")

        # Try reportlab fallback if all else fails
        if not conversion_successful:
            logger.info("Falling back to reportlab PDF generation")
            if convert_with_reportlab(output_docx, output_pdf, name, domain, start_date, end_date, gender):
                logger.info("Successfully created PDF using reportlab fallback")
                conversion_successful = True
            else:
                logger.error("All PDF conversion methods failed")

        if not conversion_successful:
            raise Exception("PDF conversion failed. All methods (docx2pdf, LibreOffice, reportlab) failed.")

        # Clean up the temporary DOCX file
        if os.path.exists(output_docx):
            os.remove(output_docx)

        print("Returning PDF")

        return output_pdf

    except Exception as e:
        raise Exception(f"Error generating certificate: {str(e)}")

@app.route('/health', methods=['GET'])
def health_check():
    """Health check endpoint"""
    health_info = {
        "status": "healthy",
        "service": "certificate-generator",
        "docx2pdf_available": DOCX2PDF_AVAILABLE,
        "reportlab_available": REPORTLAB_AVAILABLE,
        "template_exists": os.path.exists("SpectoV_Cert.docx"),
        "current_directory": os.getcwd(),
        "python_version": platform.python_version(),
        "platform": platform.system()
    }

    # Check LibreOffice availability
    libreoffice_available = False
    libreoffice_commands = ['libreoffice', 'soffice', '/usr/bin/libreoffice', '/usr/bin/soffice']
    for cmd in libreoffice_commands:
        try:
            result = subprocess.run([cmd, '--version'], capture_output=True, text=True, timeout=5)
            if result.returncode == 0:
                libreoffice_available = True
                break
        except (subprocess.TimeoutExpired, FileNotFoundError):
            continue

    health_info["libreoffice_available"] = libreoffice_available

    return jsonify(health_info)

@app.route('/generate-certificate', methods=['POST'])
def generate_certificate_api():
    """
    API endpoint to generate certificate
    Expected JSON payload:
    {
        "name": "John Doe",
        "domain": "Web Development",
        "start_date": "January 1, 2024",
        "end_date": "March 31, 2024",
        "gender": "male"  // optional, defaults to "other"
    }
    """
    try:
        # Get JSON data from request
        data = request.get_json()
        
        if not data:
            return jsonify({
                "success": False,
                "error": "No JSON data provided"
            }), 400
        
        # Extract required fields
        name = data.get('name')
        domain = data.get('domain')
        start_date = data.get('start_date')
        end_date = data.get('end_date')
        gender = data.get('gender', 'other')
        
        # Validate required fields
        if not all([name, domain, start_date, end_date]):
            return jsonify({
                "success": False,
                "error": "Missing required fields: name, domain, start_date, end_date"
            }), 400
        
        # Generate certificate
        pdf_path = generate_certificate(name, domain, start_date, end_date, gender)
        
        # Return the PDF file
        return send_file(
            pdf_path,
            as_attachment=True,
            download_name=f"certificate_{name.replace(' ', '_')}.pdf",
            mimetype='application/pdf'
        )
        
    except FileNotFoundError as e:
        return jsonify({
            "success": False,
            "error": str(e)
        }), 404
        
    except Exception as e:
        return jsonify({
            "success": False,
            "error": str(e)
        }), 500
    
    finally:
        # Clean up generated PDF file after sending
        try:
            if 'pdf_path' in locals() and os.path.exists(pdf_path):
                # Schedule cleanup after response is sent
                @app.after_request
                def cleanup_file(response):
                    try:
                        if os.path.exists(pdf_path):
                            os.remove(pdf_path)
                    except:
                        pass  # Ignore cleanup errors
                    return response
        except:
            pass  # Ignore cleanup errors

@app.errorhandler(404)
def not_found(error):
    return jsonify({
        "success": False,
        "error": "Endpoint not found"
    }), 404

@app.errorhandler(500)
def internal_error(error):
    return jsonify({
        "success": False,
        "error": "Internal server error"
    }), 500

if __name__ == '__main__':
    # Check if template file exists
    if not os.path.exists("SpectoV_Cert.docx"):
        logger.warning("Certificate template 'SpectoV_Cert.docx' not found in current directory")
        logger.info(f"Current working directory: {os.getcwd()}")
        logger.info(f"Files in current directory: {os.listdir('.')}")

    logger.info("Starting Certificate Generation Service...")

    # Get port from environment variable (for Render.com) or default to 5001
    port = int(os.environ.get('PORT', 5001))

    logger.info(f"Health check: http://localhost:{port}/health")
    logger.info(f"Generate certificate: POST http://localhost:{port}/generate-certificate")

    # Run Flask app - use 0.0.0.0 to bind to all interfaces
    app.run(host='0.0.0.0', port=port, debug=False)
