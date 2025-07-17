from docx import Document
from datetime import datetime
from docx2pdf import convert
import os
import sys
import json

def generate_certificate(name, domain, start_date, end_date, gender):
    """
    Generate a certificate with the provided details
    """
    try:
        # === Pronoun mapping ===
        pronouns = {"male": ("he", "him"), "female": ("she", "her"), "other": ("they", "them")}
        he_she, him_her = pronouns.get(gender.lower(), ("they", "them"))

        issued_date = datetime.today().strftime('%B %d, %Y')

        # === Load and fill Word document ===
        template_path = "SpectoV_Cert.docx"
        if not os.path.exists(template_path):
            raise FileNotFoundError(f"Certificate template not found: {template_path}")
            
        doc = Document(template_path)

        # Replace placeholders in all paragraphs
        for para in doc.paragraphs:
            if "{{Domain}}" in para.text:
                para.text = para.text.replace("{{Domain}}", domain)
            if "{{Start Date}}" in para.text:
                para.text = para.text.replace("{{Start Date}}", start_date)
            if "{{End Date}}" in para.text:
                para.text = para.text.replace("{{End Date}}", end_date)
            if "{{he/she/they}}" in para.text:
                para.text = para.text.replace("{{he/she/they}}", he_she)
            if "{{him/her/them}}" in para.text:
                para.text = para.text.replace("{{him/her/them}}", him_her)
            if "{{Name}}" in para.text:
                para.text = para.text.replace("{{Name}}", name)
            if "ISSUED DATE :" in para.text:
                para.text = para.text.replace("ISSUED DATE :", f"ISSUED DATE : {issued_date}")

        # Also check tables for placeholders
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if "{{Domain}}" in para.text:
                            para.text = para.text.replace("{{Domain}}", domain)
                        if "{{Start Date}}" in para.text:
                            para.text = para.text.replace("{{Start Date}}", start_date)
                        if "{{End Date}}" in para.text:
                            para.text = para.text.replace("{{End Date}}", end_date)
                        if "{{he/she/they}}" in para.text:
                            para.text = para.text.replace("{{he/she/they}}", he_she)
                        if "{{him/her/them}}" in para.text:
                            para.text = para.text.replace("{{him/her/them}}", him_her)
                        if "{{Name}}" in para.text:
                            para.text = para.text.replace("{{Name}}", name)
                        if "ISSUED DATE :" in para.text:
                            para.text = para.text.replace("ISSUED DATE :", f"ISSUED DATE : {issued_date}")

        # === Save DOCX ===
        output_docx = "Final_Certificate.docx"
        doc.save(output_docx)

        # === Convert to PDF ===
        output_pdf = "Final_Certificate.pdf"
        convert(output_docx, output_pdf)

        # Clean up the temporary DOCX file
        if os.path.exists(output_docx):
            os.remove(output_docx)

        return output_pdf

    except Exception as e:
        raise Exception(f"Error generating certificate: {str(e)}")

if __name__ == "__main__":
    try:
        # Read input from command line arguments or stdin
        if len(sys.argv) > 1:
            # Input from command line arguments (JSON string)
            input_data = json.loads(sys.argv[1])
        else:
            # Input from stdin (for interactive mode)
            input_data = json.loads(sys.stdin.read())
        
        name = input_data.get('name')
        domain = input_data.get('domain')
        start_date = input_data.get('start_date')
        end_date = input_data.get('end_date')
        gender = input_data.get('gender', 'other')
        
        if not all([name, domain, start_date, end_date]):
            raise ValueError("Missing required fields: name, domain, start_date, end_date")
        
        pdf_path = generate_certificate(name, domain, start_date, end_date, gender)
        
        # Return success response
        result = {
            "success": True,
            "pdf_path": pdf_path,
            "message": f"Certificate generated successfully: {pdf_path}"
        }
        print(json.dumps(result))
        
    except Exception as e:
        # Return error response
        result = {
            "success": False,
            "error": str(e)
        }
        print(json.dumps(result))
        sys.exit(1)
